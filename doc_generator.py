# =====================================================
# DOC GENERATOR - PROFISSIONAL
# =====================================================

from docx import Document
from datetime import datetime
import os
import re
from typing import Optional

# =====================================================
# CAMPOS EM NEGRITO
# =====================================================
CAMPOS_NEGRITO = (
    "CONTRATANTE_",
    "CONTRATADA_",
    "BANCO",
    "AGENCIA",
    "CONTA",
    "CONTA_CORRENTE",
)

# =====================================================
# NORMALIZAR TEXTO
# =====================================================
def normalizar_texto(texto: str) -> str:
    if not texto:
        return ""
    texto = texto.replace("\t", " ")
    texto = re.sub(r" {2,}", " ", texto)
    return texto.strip()

# =====================================================
# SUBSTITUIÇÃO EM PARÁGRAFOS
# =====================================================
def _substituir_em_paragrafo(paragrafo, dados: dict):
    if not paragrafo.runs:
        return

    texto_completo = "".join(run.text for run in paragrafo.runs)
    texto_original = texto_completo
    aplicar_negrito = False

    for chave, valor in dados.items():
        valor = "" if valor is None else str(valor)

        padrao = r"\{\{\s*" + re.escape(chave) + r"\s*\}\}"

        if re.search(padrao, texto_completo):
            if chave.startswith(CAMPOS_NEGRITO):
                aplicar_negrito = True

            texto_completo = re.sub(padrao, valor, texto_completo)

    if texto_completo == texto_original:
        return

    texto_completo = normalizar_texto(texto_completo)

    run_principal = paragrafo.runs[0]
    run_principal.text = texto_completo

    if aplicar_negrito:
        run_principal.bold = True

    for run in paragrafo.runs[1:]:
        run.text = ""

# =====================================================
# TABELAS
# =====================================================
def _substituir_em_tabelas(doc, dados: dict):
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_em_paragrafo(paragrafo, dados)

# =====================================================
# GERADOR PRINCIPAL (COMPATÍVEL COM APP)
# =====================================================
def gerar_documento(
    template_path: str,
    output_path: str,
    dados: dict = None
) -> str:

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    dados = dados or {}

    doc = Document(template_path)

    # Substituições
    for paragrafo in doc.paragraphs:
        _substituir_em_paragrafo(paragrafo, dados)

    _substituir_em_tabelas(doc, dados)

    # Garante pasta
    pasta = os.path.dirname(output_path)
    if pasta:
        os.makedirs(pasta, exist_ok=True)

    doc.save(output_path)

    return output_path