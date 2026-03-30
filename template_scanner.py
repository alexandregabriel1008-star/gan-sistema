from docx import Document
import re

PADROES = {
    "CONTRATANTE": [
        "CONTRATANTE_NOME",
        "CONTRATANTE_CNPJ",
        "CONTRATANTE_ENDERECO",
        "CONTRATANTE_CIDADE_UF",
        "CONTRATANTE_RESPONSAVEL",
        "REGIME_TRIBUTARIO",
        "INICIO_ATIVIDADES",
    ],
    "CONTRATADA": [
        "CONTRATADA_NOME",
        "CONTRATADA_CNPJ",
        "CONTRATADA_ENDERECO",
        "CONTRATADA_CIDADE_UF",
        "CONTRATADA_RESPONSAVEL",
        "BANCO",
        "AGENCIA",
        "CONTA_CORRENTE",
    ],
    "FINANCEIRO": [
        "VALOR",
        "PERCENTUAL",
    ],
    "GERAL": [
        "DATA_ATUALIZADA",
    ]
}


def scan_template(template_path):
    doc = Document(template_path)
    texto = []

    for p in doc.paragraphs:
        texto.append(p.text)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto.append(celula.text)

    conteudo = "\n".join(texto)

    encontrados = sorted(set(re.findall(r"\{\{(.*?)\}\}", conteudo)))

    classificados = {
        "CONTRATANTE": [],
        "CONTRATADA": [],
        "FINANCEIRO": [],
        "GERAL": [],
        "DESCONHECIDOS": [],
    }

    for var in encontrados:
        achou = False
        for grupo, lista in PADROES.items():
            if var in lista:
                classificados[grupo].append(var)
                achou = True
                break
        if not achou:
            classificados["DESCONHECIDOS"].append(var)

    return classificados
