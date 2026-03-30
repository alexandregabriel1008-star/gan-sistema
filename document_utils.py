from docx import Document
from datetime import datetime
import os
import re
from typing import Optional

# =====================================================
# CAMPOS QUE DEVEM FICAR EM NEGRITO
# =====================================================
CAMPOS_NEGRITO = (
    "CONTRATANTE_",
    "CONTRATADA_",
    "BANCO",
    "AGENCIA",
    "CONTA",
    "CONTA_CORRENTE",
)

# =====================================================
# NORMALIZA TEXTO (LEVE – NÃO QUEBRA FORMATAÇÃO)
# =====================================================
def normalizar_texto(texto: str) -> str:
    texto = texto.replace("\t", " ")
    texto = re.sub(r" {2,}", " ", texto)
    return texto.strip()


# =====================================================
# SUBSTITUI VARIÁVEIS EM PARÁGRAFOS (ROBUSTO)
# =====================================================
def _substituir_em_paragrafo(paragrafo, dados: dict):

    if not paragrafo.runs:
        return

    texto_completo = "".join(run.text for run in paragrafo.runs)
    texto_original = texto_completo
    aplicar_negrito = False

    for chave, valor in dados.items():
        valor = "" if valor is None else str(valor)

        padrao = r"\{\{\s*" + re.escape(chave) + r"\s*\}\}"

        if re.search(padrao, texto_completo):
            if chave.startswith(CAMPOS_NEGRITO):
                aplicar_negrito = True

            texto_completo = re.sub(padrao, valor, texto_completo)

    if texto_completo == texto_original:
        return

    texto_completo = normalizar_texto(texto_completo)

    # Mantém apenas o primeiro run
    run_principal = paragrafo.runs[0]
    run_principal.text = texto_completo

    if aplicar_negrito:
        run_principal.bold = True

    # Limpa runs restantes
    for run in paragrafo.runs[1:]:
        run.text = ""


# =====================================================
# SUBSTITUI VARIÁVEIS EM TABELAS
# =====================================================
def _substituir_em_tabelas(doc, dados: dict):
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_em_paragrafo(paragrafo, dados)


# =====================================================
# GERADOR PRINCIPAL
# =====================================================
def gerar_documento(
    template_path: str,
    dados: dict,
    tipo: str = "DOCUMENTO",
    cliente: Optional[str] = None,
    pasta_destino: Optional[str] = None
) -> str:

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    doc = Document(template_path)

    # Parágrafos
    for paragrafo in doc.paragraphs:
        _substituir_em_paragrafo(paragrafo, dados)

    # Tabelas
    _substituir_em_tabelas(doc, dados)

    # Pasta destino
    if not pasta_destino:
        pasta_destino = "output"

    os.makedirs(pasta_destino, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cliente_slug = cliente.replace(" ", "_") if cliente else "GERAL"
    tipo_slug = tipo.upper()

    nome_arquivo = f"{tipo_slug}_{cliente_slug}_{timestamp}.docx"
    caminho = os.path.join(pasta_destino, nome_arquivo)

    doc.save(caminho)
    return caminho
